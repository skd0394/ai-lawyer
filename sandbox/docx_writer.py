"""Markdown -> legal-styled .docx. Runs INSIDE the sandbox.

The model writes MARKDOWN. This module owns presentation.

Why not have the model emit OOXML directly:
  - output tokens: a 4k-token contract becomes 20k+ of XML, billed at ~5x
    input rates, and output is the one thing caching cannot help with
  - reliability: models lose track of closing tags on long documents and
    produce files Word refuses to open
  - consistency: a deterministic converter makes every document look the
    same, rather than however the model felt that turn

Deliberately NOT implemented: automatic clause numbering. python-docx list
numbering means editing numbering.xml and is a time sink. The model writes
"## 5. Termination" and the number is plain text — which is also what a
drafter wants, since legal numbering is semantic (5. vs 5.1 vs (e)).
"""

from __future__ import annotations

import os
import re

# [PARTY A NAME], [DATE OF EXECUTION], [SECURITY DEPOSIT AMOUNT]
PLACEHOLDER_RE = re.compile(r"\[([A-Z][A-Z0-9 _/\-.,'&]{1,60})\]")

# Bump on every change to this file. Returned with every generated
# document so you can tell at a glance whether the sandbox is running the
# code you just wrote — stale workers otherwise look exactly like bugs.
WRITER_VERSION = 2

BODY_FONT = "Times New Roman"
BODY_SIZE = 11


def _page_number_field(paragraph) -> None:
    """Word field codes for 'Page X of Y'. python-docx has no API for this,
    so we build the XML by hand."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def field(code: str):
        run = paragraph.add_run()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = code
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.append(begin)
        run._r.append(instr)
        run._r.append(end)

    paragraph.add_run("Page ")
    field("PAGE")
    paragraph.add_run(" of ")
    field("NUMPAGES")


def _add_runs(paragraph, text: str) -> None:
    """Inline **bold**, *italic*, and [PLACEHOLDERS].

    Placeholders are bolded so they are impossible to miss when the user
    opens the document — the whole point is that they get filled in.
    """
    token = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|\[[A-Z][A-Z0-9 _/\-.,'&]{1,60}\])")
    for part in token.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            paragraph.add_run(part[1:-1]).italic = True
        elif PLACEHOLDER_RE.fullmatch(part):
            r = paragraph.add_run(part)
            r.bold = True
        else:
            paragraph.add_run(part)


def _parse_blocks(md: str) -> list[dict]:
    """Minimal block parser: headings, tables, lists, rules, paragraphs."""
    blocks: list[dict] = []
    lines = md.replace("\r\n", "\n").split("\n")
    i = 0

    while i < len(lines):
        line = lines[i].rstrip()
        st = line.strip()

        if not st:
            i += 1
            continue

        if st.startswith("#"):
            level = len(st) - len(st.lstrip("#"))
            blocks.append({"t": "h", "level": min(level, 4),
                           "text": st.lstrip("#").strip()})
            i += 1
            continue

        if re.fullmatch(r"(-{3,}|\*{3,}|_{3,})", st):
            blocks.append({"t": "rule"})
            i += 1
            continue

        # signature line: a run of underscores
        if re.fullmatch(r"_{5,}.*", st):
            blocks.append({"t": "sig", "text": st})
            i += 1
            continue

        if st.startswith("|") and i + 1 < len(lines) and \
                re.match(r"^\|[\s:|-]+\|?$", lines[i + 1].strip()):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not re.match(r"^[\s:|-]+$", "|".join(cells)):
                    rows.append(cells)
                i += 1
            if rows:
                blocks.append({"t": "table", "rows": rows})
            continue

        m = re.match(r"^(\s*)([-*+]|\d+[.)])\s+(.*)$", line)
        if m:
            items = []
            ordered = not m.group(2) in ("-", "*", "+")
            while i < len(lines):
                mm = re.match(r"^(\s*)([-*+]|\d+[.)])\s+(.*)$", lines[i])
                if not mm:
                    break
                items.append({"text": mm.group(3).strip(),
                              "indent": len(mm.group(1)) // 2})
                i += 1
            blocks.append({"t": "list", "ordered": ordered, "items": items})
            continue

        para = [st]
        i += 1
        while i < len(lines) and lines[i].strip() and \
                not lines[i].strip().startswith(("#", "|", "-", "*", "_")):
            para.append(lines[i].strip())
            i += 1
        blocks.append({"t": "p", "text": " ".join(para)})

    return blocks


def markdown_to_docx(md: str, title: str, out_path: str,
                     subtitle: str = "") -> dict:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, Inches, RGBColor

    doc = Document()

    # Base style. Times New Roman is the convention for legal documents;
    # justified body text is what makes it look like a contract rather
    # than a memo.
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(BODY_SIZE)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    # Title block. add_heading(level=0) applies Word's "Title" STYLE, which
    # matters because extract.py reads styles back to reconstruct markdown.
    # A bold run in a Normal paragraph looks identical on screen but is
    # invisible to the round trip.
    tp = doc.add_heading(title.upper(), level=0)
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for tr in tp.runs:
        tr.bold = True
        tr.font.size = Pt(14)
        tr.font.name = BODY_FONT
        tr.font.color.rgb = RGBColor(0, 0, 0)   # Word's default is blue
    if subtitle:
        sp = doc.add_paragraph()
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = sp.add_run(subtitle)
        sr.italic = True
        sr.font.size = Pt(10)
    doc.add_paragraph()

    first_heading_skipped = False
    for b in _parse_blocks(md):
        if b["t"] == "h":
            # The model usually repeats the title as an H1; don't print it
            # twice under the title block.
            if (not first_heading_skipped and b["level"] == 1
                    and b["text"].strip().lower() == title.strip().lower()):
                first_heading_skipped = True
                continue
            # ⭐ add_heading, NOT add_paragraph + bold run.
            # Heading STYLES are what make a document machine-readable:
            # extract.py maps "Heading N" back to markdown "#" levels, which
            # is what powers read_document(mode="section") and
            # edit_document. Bold text in a Normal paragraph looks the same
            # to a human and is invisible to the round trip — which silently
            # broke draft-then-revise on every document the agent produced.
            p = doc.add_heading(b["text"], level=b["level"])
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(12 if b["level"] <= 2 else 11)
                r.font.name = BODY_FONT
                r.font.color.rgb = RGBColor(0, 0, 0)

        elif b["t"] == "p":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _add_runs(p, b["text"])

        elif b["t"] == "list":
            for item in b["items"]:
                p = doc.add_paragraph(
                    style="List Number" if b["ordered"] else "List Bullet")
                p.paragraph_format.left_indent = Inches(
                    0.35 + 0.3 * item["indent"])
                _add_runs(p, item["text"])

        elif b["t"] == "table":
            rows = b["rows"]
            t = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
            t.style = "Table Grid"
            for ri, row in enumerate(rows):
                for ci, cell in enumerate(row):
                    para = t.rows[ri].cells[ci].paragraphs[0]
                    _add_runs(para, cell)
                    if ri == 0:
                        for run in para.runs:
                            run.bold = True
            doc.add_paragraph()

        elif b["t"] == "rule":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run("* * *")

        elif b["t"] == "sig":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.add_run(b["text"])

    # Footer page numbers — small thing, large effect on how finished it looks
    footer_p = doc.sections[0].footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _page_number_field(footer_p)
    for run in footer_p.runs:
        run.font.size = Pt(9)
        run.font.name = BODY_FONT

    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)

    placeholders = sorted(set(PLACEHOLDER_RE.findall(md)))

    # Which paragraph styles actually landed in the file. If "Heading 1"
    # is absent, headings were written as plain bold text and the document
    # will NOT round-trip for section reads or edits.
    styles_used = sorted({p.style.name for p in doc.paragraphs})

    return {
        "writer_version": WRITER_VERSION,
        "styles_used": styles_used,
        "has_heading_styles": any(s.lower().startswith(("heading", "title"))
                                  for s in styles_used),
        "path": out_path,
        "bytes": os.path.getsize(out_path),
        "word_count": len(re.sub(r"[#*|_-]", " ", md).split()),
        # Returned so the agent can tell the user what is still missing.
        # The spec: missing details must never BLOCK drafting.
        "placeholders": [f"[{p}]" for p in placeholders],
        "placeholder_count": len(placeholders),
    }


# ── section-scoped editing ────────────────────────────────────────────────
# Why sections rather than whole-document rewrites: regenerating a
# 3,764-word agreement to change one clause costs ~5,000 OUTPUT tokens
# (~$0.075 at Sonnet rates), and output is the one thing prompt caching
# cannot help with. A section edit is ~200 tokens. Same principle as
# extract-then-discard, applied to the output side.

def find_section_span(md: str, name: str) -> tuple[int, int, int] | None:
    """Character span of a section: its heading through to the next heading
    of the same or higher level. Returns (start, end, level)."""
    want = (name or "").strip().lstrip("#").strip().lower()
    if not want:
        return None

    heads = []
    for m in re.finditer(r"^(#{1,6})[ \t]+(.+)$", md, re.M):
        heads.append((m.start(), m.end(), len(m.group(1)), m.group(2).strip()))

    for idx, (start, _end, level, text) in enumerate(heads):
        norm = text.lower()
        # Match on containment either way, so "Termination" finds
        # "5. Termination" and "5. Termination" finds "Termination".
        if want in norm or norm in want:
            stop = len(md)
            for nstart, _ne, nlevel, _nt in heads[idx + 1:]:
                if nlevel <= level:
                    stop = nstart
                    break
            return (start, stop, level)
    return None


def list_sections(md: str) -> list[str]:
    return [m.group(2).strip()
            for m in re.finditer(r"^(#{1,6})[ \t]+(.+)$", md, re.M)]


def apply_edits(md: str, edits: list[dict]) -> tuple[str, list, list]:
    """Apply section-scoped edits in order. Returns (new_md, applied, failed).

    Failures are RETURNED, not raised — a bad section name should come back
    to the model as something it can correct, along with the list of
    sections that do exist.
    """
    applied, failed = [], []

    for e in edits or []:
        action = (e.get("action") or "replace").lower()
        section = e.get("section") or ""
        content = (e.get("content") or "").rstrip()

        if action == "append_document":
            md = md.rstrip() + "\n\n" + content + "\n"
            applied.append({"action": action, "section": "(end of document)"})
            continue

        span = find_section_span(md, section)
        if span is None:
            failed.append({"action": action, "section": section,
                           "reason": "no matching section",
                           "available": list_sections(md)[:30]})
            continue

        start, end, _level = span
        before, target, after = md[:start], md[start:end], md[end:]

        if action == "replace":
            md = before + content + "\n\n" + after
        elif action in ("append_to_section", "insert_after"):
            md = before + target.rstrip() + "\n\n" + content + "\n\n" + after
        elif action == "delete":
            md = before + after
        else:
            failed.append({"action": action, "section": section,
                           "reason": f"unknown action '{action}'",
                           "available": ["replace", "append_to_section",
                                         "insert_after", "delete",
                                         "append_document"]})
            continue

        applied.append({"action": action, "section": section,
                        "chars_before": len(target), "chars_after": len(content)})

    return md, applied, failed